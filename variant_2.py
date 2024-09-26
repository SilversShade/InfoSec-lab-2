import os
import sys

import docx
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from docx.text.run import Run


def insert_secret_message(txt_filename: str, docx_filename: str, pt_amplifier: float) -> str:
    with open(txt_filename, "r", encoding="utf-8") as file:
        secret_message = file.read()

    document = docx.Document(docx_filename)
    docx_content = ""
    font_size = document.paragraphs[0].style.font.size.pt
    for paragraph in document.paragraphs:
        docx_content += paragraph.text + '\n'

    indices = []
    index = 0
    docx_content_lower = docx_content.lower()
    for letter in secret_message:
        index = docx_content_lower.find(letter.lower(), index + 1)

        if index == -1:
            sys.exit("Could not find an appropriate letter in the initial docx file")
        indices.append(index)

    docx_result = docx.Document()
    docx_result_text: Paragraph = docx_result.add_paragraph()
    run: Run = docx_result_text.add_run(docx_content[:indices[0]])
    run.font.size = Pt(font_size)
    for index in range(len(indices)):
        run: Run = docx_result_text.add_run(docx_content[indices[index]])
        run.font.size = Pt(font_size + pt_amplifier)
        if index == (len(indices) - 1):
            run = docx_result_text.add_run(docx_content[indices[index] + 1:])
        else:
            run = docx_result_text.add_run(docx_content[indices[index] + 1:indices[index + 1]])
        run.font.size = Pt(font_size)

    split = os.path.splitext(docx_filename)
    result_filename = split[0] + "_result" + split[1]

    docx_result.save(result_filename)
    return result_filename


def read_secret_message(docx_filename: str):
    document = docx.Document(docx_filename)

    pt_to_text = {}
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            font_size = run.font.size.pt
            if font_size not in pt_to_text:
                pt_to_text[font_size] = ""
            pt_to_text[font_size] += run.text


    print(pt_to_text[sorted(pt_to_text)[-1]])


def main(txt_filename: str, docx_filename: str, pt_amplifier: float):
    docx_result_filename = insert_secret_message(txt_filename, docx_filename, pt_amplifier)
    read_secret_message(docx_result_filename)


main("txt.txt", "docx.docx", 10)
