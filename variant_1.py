import sys
import os

import docx
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.shared import RGBColor


def insert_secret_message(txt_filename: str, docx_filename: str, red_value: int) -> str:
    with open(txt_filename, "r", encoding="utf-8") as file:
        secret_message = file.read()

    document = docx.Document(docx_filename)

    docx_content = ""
    for paragraph in document.paragraphs:
        docx_content += paragraph.text + '\n'

    indices = []
    index = 0
    docx_content_lower = docx_content.lower()
    for letter in secret_message:
        index = docx_content_lower.find(letter.lower(), index + 1)

        if index == -1:
            sys.exit("Could not find an appropriate letter in the initial docx file")
        indices.append(index)

    docx_result = docx.Document()
    docx_result_text: Paragraph = docx_result.add_paragraph(docx_content[:indices[0]])
    for index in range(len(indices)):
        run: Run = docx_result_text.add_run(docx_content[indices[index]])
        run.font.color.rgb = docx.shared.RGBColor(red_value, 0, 0)
        if index == (len(indices) - 1):
            docx_result_text.add_run(docx_content[indices[index] + 1:])
        else:
            docx_result_text.add_run(docx_content[indices[index] + 1:indices[index + 1]])

    split = os.path.splitext(docx_filename)
    result_filename = split[0] + "_result" + split[1]

    docx_result.save(result_filename)
    return result_filename


def read_secret_message(docx_filename: str, red_value: int):
    document = docx.Document(docx_filename)

    result = ""
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            color = run.font.color.rgb
            if color is not None and color == RGBColor(red_value, 0, 0):
                result += run.text

    print(result)


def main(txt_filename: str, docx_filename: str, red_value: int):
    docx_result_filename = insert_secret_message(txt_filename, docx_filename, red_value)
    read_secret_message(docx_result_filename, red_value)


main("txt.txt", "docx.docx", 255)
